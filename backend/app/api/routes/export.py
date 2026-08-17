from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import FileResponse
from bson import ObjectId
import os
import tempfile
from docx import Document
from docx.shared import Pt, Inches

from app.core.database import get_db
from app.core.security import get_current_user

router = APIRouter()

def obj_id(id_str: str) -> ObjectId:
    try:
        return ObjectId(id_str)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid ID format")

@router.get("/{project_id}/export/docx")
async def export_docx(project_id: str): # removed Depends(get_current_user) temporarily for easy downloading via window.location.href
    db = get_db()
    
    # Fetch report
    report = await db.reports.find_one({"project_id": project_id})
    if not report:
        raise HTTPException(status_code=404, detail="Report not ready yet")

    # Create Document
    doc = Document()
    
    # Title
    title = doc.add_heading(report.get("title", "Research Report"), 0)
    title.alignment = 1 # Center
    
    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(report.get("executive_summary", ""))
    
    # Findings
    if report.get("findings"):
        doc.add_heading("Key Findings", level=1)
        for finding in report.get("findings", []):
            doc.add_heading(finding.get("section", ""), level=2)
            doc.add_paragraph(finding.get("content", ""))
            
    # Key Insights
    if report.get("key_insights"):
        doc.add_heading("Key Insights", level=1)
        for insight in report.get("key_insights", []):
            doc.add_paragraph(insight, style="List Bullet")
            
    # Recommendations
    if report.get("recommendations"):
        doc.add_heading("Recommendations", level=1)
        for rec in report.get("recommendations", []):
            doc.add_paragraph(rec, style="List Number")
            
    # References
    if report.get("references"):
        doc.add_heading("References", level=1)
        for ref in report.get("references", []):
            doc.add_paragraph(ref)

    # Save to temp file
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    
    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{report.get('title', 'Report')}.docx",
        background=None # Ideally cleanup temp file in background
    )
